# -*- coding: utf-8 -*-
"""
同类开源方案调研报告生成脚本
课题：实体链接与知识对齐智能体（课题10）
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shading_elm)


def add_table_row(table, cells, bold=False, shading=None):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if bold:
            run.bold = True
        if shading:
            set_cell_shading(cell, shading)
    return row


def create_report():
    doc = Document()

    # ========== 设置默认样式 ==========
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置段落间距
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # ========== 封面 ==========
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('同类开源方案调研报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('课题10：实体链接与知识对齐智能体')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    info_lines = [
        f'调研日期：{datetime.date.today().strftime("%Y年%m月%d日")}',
        '项目名称：ELAGENT — 实体链接与知识对齐智能体',
        '所属课题：数据治理流水线 · 专项能力智能体（课题10）',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ========== 目录页 ==========
    toc_title = doc.add_heading('目  录', level=1)
    toc_items = [
        ('一、课题题目与定位', 3),
        ('二、同类开源方案总览', 3),
        ('三、主要开源工具详细分析', 4),
        ('  3.1 BLINK（Facebook AI）', 4),
        ('  3.2 ReFinED（Amazon/Google）', 5),
        ('  3.3 GENRE（Meta AI）', 5),
        ('  3.4 REL（Radboud University）', 6),
        ('  3.5 DBpedia Spotlight', 6),
        ('  3.6 其他重要工具', 7),
        ('四、最新研究进展（2023–2026）', 7),
        ('  4.1 LLM驱动的实体链接', 7),
        ('  4.2 零样本/少样本实体链接', 8),
        ('  4.3 NIL检测研究进展', 8),
        ('  4.4 中文实体链接研究', 9),
        ('五、本智能体与通用/开源方案的差异与攻关点', 9),
        ('六、结论与建议', 11),
        ('参考文献', 12),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}')
        run.font.size = Pt(12)
        if not item.startswith('  '):
            run.bold = True

    doc.add_page_break()

    # ========== 一、课题题目与定位 ==========
    doc.add_heading('一、课题题目与定位', level=1)

    doc.add_heading('1.1 课题题目', level=2)
    p = doc.add_paragraph()
    run = p.add_run('实体链接与知识对齐智能体（课题10）')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_heading('1.2 课题定位', level=2)
    doc.add_paragraph(
        '本课题属于数据治理流水线中的"专项能力智能体"，聚焦于实体识别（NER）之后的'
        '"链接、消歧与标准化"环节。在完整的数据治理流水线中，课题10承担以下关键角色：'
    )
    bullets = [
        '被课题4（全域术语归一智能体）调用，完成实体消歧与标准化',
        '被课题3（数据清洗智能体）调用，用于人名/机构脱敏定位',
        '作为独立服务，对外暴露标准接口供流水线以注册制热插拔集成',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('1.3 研究目标', level=2)
    doc.add_paragraph(
        '在已给定文本与实体指称（mention）的前提下，把每个指称链接到知识库中的标准实体'
        '（标准全称 + 唯一ID），完成消歧、别名/简称/曾用名标准化、NIL检测与共指消解，'
        '并保留链接依据以供追溯。'
    )

    # 技术指标表
    doc.add_heading('1.4 验收指标', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ['指标项', '目标值', '说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, '2E75B6')
        run.font.color.rgb = RGBColor(255, 255, 255)

    metrics = [
        ('链接准确率', '≥ 85%', '指称→标准实体链接准确率'),
        ('消歧准确率', '≥ 85%', '同名异指消歧准确率'),
        ('NIL检测F1', '≥ 0.80', '知识库中不存在实体的检测'),
        ('别名标准化召回率', '≥ 85%', '别名/简称/曾用名映射'),
        ('共指消解准确率', '≥ 80%', '代词/指代回链（按需启用）'),
        ('可追溯', '100%', '链接结果均留存依据'),
    ]
    for row_data in metrics:
        add_table_row(table, row_data)

    doc.add_page_break()

    # ========== 二、同类开源方案总览 ==========
    doc.add_heading('二、同类开源方案总览', level=1)

    doc.add_paragraph(
        '实体链接（Entity Linking, EL）是自然语言处理的核心任务之一，旨在将文本中的实体指称'
        '（mention）链接到知识库中的标准实体。近年来，该领域涌现出大量开源工具和研究成果。'
        '下表对主流开源方案进行了总览比较：'
    )

    # 总览表
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['工具名称', '开发机构', '核心方法', '开源地址', '特点', '局限性']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, '2E75B6')
        run.font.color.rgb = RGBColor(255, 255, 255)

    tools = [
        ('BLINK', 'Facebook AI', 'Bi-encoder + Cross-encoder',
         'github.com/facebookresearch/BLINK',
         '快速、多语言支持、鲁棒性强',
         '绑定静态Wikipedia数据；hard negatives未实现'),
        ('ReFinED', 'Amazon/Google', 'Transformer + 细粒度类型 + 描述',
         'github.com/amazon-science/ReFinED',
         '轻量高效、零样本能力强、支持Wikidata 9000万实体',
         '需要微调以泛化到非Wikipedia域'),
        ('GENRE', 'Meta AI', '自回归生成式实体检索',
         'github.com/facebookresearch/GENRE',
         '创新的生成式方法、无需向量索引',
         '推理速度较慢；零样本性能有限'),
        ('REL', 'Radboud Univ.', 'Wikipedia图 + 统计消歧',
         'github.com/informagi/REL',
         '高ER F1；成熟的API',
         '容易错误链接NIL提及'),
        ('DBpedia Spotlight', 'DBpedia', '向量空间 + 余弦相似度',
         'github.com/dbpedia-spotlight',
         '最早的语义方法之一、多语言支持',
         '整体性能最弱；消歧能力不足'),
        ('TAGME', 'Univ. of Pisa', 'Wikipedia链接图 + 消歧',
         'TagMe API',
         '消歧精度高；轻量级',
         'ER精度低；对非命名实体预测过多'),
        ('FELA', '2025年新作', '模块化LLM + 量化',
         'CEUR-WS 2025',
         '资源高效、模块化、多KB适配',
         '新兴方案，生态不成熟'),
        ('LLMAEL', '清华KEG', 'LLM上下文增强 + ReFinED',
         'github.com/THU-KEG/LLMAEL',
         'AIDA 92.38%准确率；即插即用',
         '依赖LLM推理开销'),
    ]
    for row_data in tools:
        add_table_row(table, row_data)

    doc.add_page_break()

    # ========== 三、主要开源工具详细分析 ==========
    doc.add_heading('三、主要开源工具详细分析', level=1)

    # 3.1 BLINK
    doc.add_heading('3.1 BLINK（Facebook AI）', level=2)
    doc.add_paragraph(
        'BLINK是Facebook AI于2020年发布的实体链接系统，采用两阶段架构：'
        'Bi-encoder快速召回候选实体，Cross-encoder精细排序。'
    )
    doc.add_paragraph('核心特点：', style='List Bullet')
    bullets = [
        'Bi-encoder：将mention context和entity description分别编码，通过内积计算相似度',
        'Cross-encoder：将mention和candidate拼接输入BERT，进行精细排序',
        '在AIDA-CoNLL数据集上达到约86%的准确率',
        '支持零样本实体链接（ZESHEL基准）',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet 2')

    doc.add_paragraph('局限性：')
    doc.add_paragraph(
        'BLINK的GitHub实现存在已知问题（hard negatives sampling未实现），'
        '且绑定静态Wikipedia数据，无法动态更新知识库。在EMNLP 2023的公平评测中，'
        'BLINK的整体F1约为68.0%，低于ReFinED的73.3%。'
    )

    # 3.2 ReFinED
    doc.add_heading('3.2 ReFinED（Amazon/Google）', level=2)
    doc.add_paragraph(
        'ReFinED是目前综合性能最优的开源实体链接系统之一，由Amazon Science于2022年发布。'
        '它采用Transformer编码器，结合细粒度实体类型和实体描述进行消歧。'
    )
    doc.add_paragraph('核心优势：')
    bullets = [
        '端到端：同时完成mention detection、fine-grained entity typing和disambiguation',
        '高效推理：单次前向传播完成所有子任务，比BLINK快约6倍',
        '零样本能力：支持Wikidata 9000万+实体，包括零样本实体',
        '在8个标准数据集上平均F1达到89.4%，超越SOTA 3.7个点',
        '支持自定义实体添加（通过additional_entities_file参数）',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph('在EMNLP 2023公平评测中的表现：')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['指标', 'ReFinED', 'BLINK', 'GENRE']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, 'D5E8F0')
    eval_data = [
        ('整体准确率', '73.3%', '46.3%', '64.6%'),
        ('ER F1', '82.7%', '74.0%', '74.2%'),
        ('消歧准确率', '89.2%', '63.8%', '87.4%'),
    ]
    for row_data in eval_data:
        add_table_row(table, row_data)

    # 3.3 GENRE
    doc.add_heading('3.3 GENRE（Meta AI）', level=2)
    doc.add_paragraph(
        'GENRE（Generative Entity Retrieval）是Meta AI提出的创新性生成式实体链接方法，'
        '将实体检索任务转化为自回归序列生成问题。模型从左到右、逐token生成目标实体名称。'
    )
    doc.add_paragraph('创新点：')
    bullets = [
        '自回归生成：无需大规模向量索引，直接生成实体名称',
        '跨编码：有效编码mention和entity的交互信息',
        'mGENRE：多语言扩展版本，支持跨语言实体链接',
        '在AIDA-CoNLL上达到93.3%的准确率（使用训练数据时）',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph(
        '局限性：推理速度较慢（自回归生成），零样本性能有限（ZESHEL基准约47%）。'
    )

    # 3.4 REL
    doc.add_heading('3.4 REL（Radboud University）', level=2)
    doc.add_paragraph(
        'REL是一个成熟的实体链接框架，基于Wikipedia图结构和统计消歧方法。'
        '它提供完整的API服务，支持多种实验配置。'
    )
    doc.add_paragraph(
        '在公平评测中，REL的ER F1高达83.0%（最高），但容易错误链接NIL提及，'
        '导致整体准确率为67.7%。其可复现性评为"very good"。'
    )

    # 3.5 DBpedia Spotlight
    doc.add_heading('3.5 DBpedia Spotlight', level=2)
    doc.add_paragraph(
        'DBpedia Spotlight是最早的语义标注工具之一（2011年），基于DBpedia知识库，'
        '使用向量空间模型和余弦相似度进行实体链接。支持多语言（英语、德语、法语、意大利语等）。'
    )
    doc.add_paragraph(
        '在公平评测中，DBpedia Spotlight是所有系统中表现最弱的，'
        '主要因为ER组件性能不足，消歧结果也不理想。但它作为早期开创性工作，'
        '对后续研究有重要启发意义。'
    )

    # 3.6 其他重要工具
    doc.add_heading('3.6 其他重要工具', level=2)

    doc.add_paragraph('ELEVATE-ID / FELA（2025）：', style='List Bullet')
    doc.add_paragraph(
        '2025年最新的模块化实体链接框架，基于紧凑的开源LLM（如UniNER量化模型），'
        '包含NER、ED和Reranking三个独立模块。FELA在RSS-500上达到91.39%的检索率，'
        '超过ReFinED的89.40%。代表了实体链接向轻量化、模块化发展的趋势。'
    )

    doc.add_paragraph('LELA（2025）：', style='List Bullet')
    doc.add_paragraph(
        '基于LLM的零样本实体链接方法，使用Qwen3-Embedding-4B进行密集检索，'
        'Qwen3-Reranker-4B进行重排序。在ZESHEL基准上达到83.11%的宏准确率，'
        '大幅超越BLINK的74.27%和GENRE的47.09%。'
    )

    doc.add_page_break()

    # ========== 四、最新研究进展 ==========
    doc.add_heading('四、最新研究进展（2023–2026）', level=1)

    doc.add_heading('4.1 LLM驱动的实体链接', level=2)
    doc.add_paragraph(
        '大语言模型（LLM）的兴起为实体链接带来了范式性变革。2024–2025年的研究主要集中在以下几个方向：'
    )

    doc.add_paragraph('（1）LLM作为上下文增强器（LLMAEL, CIKM 2025）', style='List Bullet')
    doc.add_paragraph(
        '清华KEG提出的LLMAEL框架，利用LLM生成mention-centered描述作为额外输入，'
        '保留传统EL模型执行链接。在6个标准数据集上，LLMAEL x ReFinED达到85.76%的平均准确率，'
        '微调后达到86.67%，刷新SOTA。'
    )

    doc.add_paragraph('（2）LLM作为链接代理（EL Agent, arXiv 2025）', style='List Bullet')
    doc.add_paragraph(
        '模拟人类认知流程的EL Agent：识别mention → 使用搜索工具（Wikidata、Google KG）'
        '查找候选 → 基于上下文消歧。该方法是检索器无关的（retriever-agnostic），'
        '展示了LLM在实体链接中的代理能力。'
    )

    doc.add_paragraph('（3）自适应路由与定向推理（ARTER, EMNLP 2025）', style='List Bullet')
    doc.add_paragraph(
        'Apple提出的ARTER框架，计算互补信号（embedding + LLM）将mention分为easy和hard cases，'
        '分别用低计算量链接器和LLM推理处理，在性能和效率之间取得平衡。'
    )

    doc.add_paragraph('（4）ChatEL（2024）', style='List Bullet')
    doc.add_paragraph(
        '使用decoder-only LLM（如GPT-4、LLaMA）通过提示词进行候选重排序，'
        '在零样本设置下展示了LLM的消歧能力。'
    )

    doc.add_heading('4.2 零样本/少样本实体链接', level=2)
    doc.add_paragraph(
        '零样本实体链接是近年来的研究热点，特别是在缺乏标注数据的行业领域。关键进展包括：'
    )
    bullets = [
        'ZESHEL基准：由Logeswaran等人（2019）提出，包含4个领域（Forgotten Realms、Lego、Star Trek、YuGiOh），是零样本EL的标准评测基准',
        'GenDecider（NAACL 2024）：集成"none of the candidates"判断，在ZESHEL上达到82.75%宏准确率',
        'LELA（2025）：使用开源LLM（Magistral-Small-2509）实现真正的零样本EL，在ZESHEL上达到83.11%',
        'FELA（2025）：模块化框架，使用量化模型减少计算开销，适合资源受限场景',
        'Hansel基准：中文少样本/零样本实体链接基准数据集',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('4.3 NIL检测研究进展', level=2)
    doc.add_paragraph(
        'NIL检测（识别知识库中不存在的实体）是实体链接的关键挑战之一。'
        '最新研究进展包括：'
    )
    bullets = [
        '提示链策略（Prompt Chaining）：先判断mention是否有精确匹配，无匹配则分类为NIL；'
        '再从候选中选择最相似实体，无法确定时仍归为NIL（EACL 2026）',
        '不确定性估计：利用LLM的logits和语义熵估计NIL概率（CEUR-WS 2025）',
        'GenDecider：专门训练的"none of the candidates"判断器，提升NIL检测精度',
        '多信号融合：结合名称匹配度、上下文相似度、类型一致性等多个信号进行NIL判定',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('4.4 中文实体链接研究', level=2)
    doc.add_paragraph(
        '中文实体链接面临独特挑战：口语化严重、短文本信息少、实体多歧义。'
        '近期重要研究包括：'
    )
    bullets = [
        '基于大语言模型的中文实体链接实证研究（自动化学报, 2025）：系统评估了LLM在中文EL上的表现',
        '知识库标记预训练孪生网络（CSA, 2022）：使用知识库标记的BERT进行实体识别，'
        'BERT-SiameseFNN进行实体消歧',
        'CCKS 2024知识融合：南京大学孙泽群团队提出大模型时代的知识融合方法，'
        '包括ZeroEA（零样本实体对齐）和DERA（密集实体检索增强）',
        '多模态实体对齐（自动化学报, 2024）：自适应特征融合方法处理多模态知识图谱对齐',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ========== 五、本智能体与通用/开源方案的差异与攻关点 ==========
    doc.add_heading('五、本智能体与通用/开源方案的差异与攻关点', level=1)

    doc.add_paragraph(
        '基于对主流开源方案和最新研究的全面调研，本智能体（ELAGENT）相较通用/开源方案'
        '在以下方面具有明确的差异与攻关点，这也是"为什么值得研究"的核心理由：'
    )

    doc.add_heading('5.1 差异对比总表', level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['维度', '通用/开源方案', 'ELAGENT（本智能体）', '攻关价值']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, '2E75B6')
        run.font.color.rgb = RGBColor(255, 255, 255)

    diff_data = [
        ('设计理念', '自主规划、开放式探索', '弱自主规划、强流程约束、全程可追溯',
         '数据治理要求结果可复现、过程可审计、改动可回滚'),
        ('知识库适配', '绑定Wikipedia/Wikidata，难以适配行业知识库',
         '支持自定义知识库、别名词典、上下位关系',
         '企业数据治理需要行业专属知识库'),
        ('NIL检测', '多数方案NIL检测能力弱（REL容易错误链接NIL）',
         '多信号融合NIL检测，专门优化',
         '现实知识库不完备，NIL检测是刚需'),
        ('可追溯性', '仅输出链接结果，无依据留痕',
         '原值→新值→依据，支持回放与回滚',
         '数据治理场景的可审计硬性要求'),
        ('实现策略', '重度依赖深度学习/LLM',
         '能用规则不用LLM：BM25、别名匹配、NIL阈值均为纯规则',
         '降低算力成本，提升可解释性'),
        ('中文支持', '多为英文为主，中文支持有限',
         '原生中文支持，中文分词+中文知识库',
         '面向国内企业数据治理场景'),
        ('服务化程度', '多为研究代码，缺乏生产级API',
         'FastAPI标准接口，可独立部署、热插拔集成',
         '流水线集成的工程化需求'),
        ('评测体系', '使用通用基准（AIDA-CoNLL等）',
         '自建评测集+回归对比+失败案例分析',
         '贴近实际治理场景的评测'),
    ]
    for row_data in diff_data:
        add_table_row(table, row_data)

    doc.add_heading('5.2 核心攻关点详解', level=2)

    doc.add_paragraph('攻关点1：行业知识库适配与动态更新', style='List Bullet')
    doc.add_paragraph(
        '通用方案（BLINK、ReFinED等）绑定了Wikipedia/Wikidata，无法直接适配企业自有的行业知识库。'
        '本智能体需要解决：如何在不重新训练模型的情况下，支持任意结构的知识库接入；'
        '如何实现知识库的动态更新（新增实体、修改别名）而不影响已有链接结果。'
    )

    doc.add_paragraph('攻关点2：低资源行业的少样本迁移', style='List Bullet')
    doc.add_paragraph(
        '行业领域往往缺乏标注数据，通用方案在零样本设置下性能大幅下降'
        '（如GENRE在ZESHEL上仅47%）。本智能体需要探索：基于规则+检索的混合方法，'
        '减少对标注数据的依赖；利用别名词典和上下位关系进行知识增强。'
    )

    doc.add_paragraph('攻关点3：NIL检测的鲁棒性', style='List Bullet')
    doc.add_paragraph(
        '现实知识库往往不完备，NIL检测是实体链接的关键瓶颈。'
        '在EMNLP 2023评测中，多数系统在NIL检测上表现不佳。'
        '本智能体采用多信号融合策略（名称匹配度、上下文相似度、类型一致性、阈值判定），'
        '需要在精确率和召回率之间取得平衡。'
    )

    doc.add_paragraph('攻关点4：全程可追溯的链接依据留痕', style='List Bullet')
    doc.add_paragraph(
        '这是本智能体与所有通用方案的根本区别。通用方案仅输出"mention → entity"的映射，'
        '不提供链接依据。本智能体要求对每一次数据改动留痕（原值→新值→依据），'
        '支持回放与回滚，契合数据治理对可审计的硬性要求。'
    )

    doc.add_paragraph('攻关点5：能用规则不用LLM的效率优化', style='List Bullet')
    doc.add_paragraph(
        '最新研究趋势是用LLM做实体链接（如ARTER、ChatEL），但LLM推理成本高、'
        '可解释性差。本智能体遵循"能用确定性规则解决就不消耗大模型算力"的原则，'
        'BM25全文检索、别名精确匹配、NIL阈值判定均为纯规则实现，'
        '仅在复杂消歧场景保留LLM作为兜底。'
    )

    doc.add_page_break()

    # ========== 六、结论与建议 ==========
    doc.add_heading('六、结论与建议', level=1)

    doc.add_heading('6.1 调研结论', level=2)
    doc.add_paragraph(
        '通过对BLINK、ReFinED、GENRE、REL、DBpedia Spotlight、FELA、LLMAEL、LELA等'
        '主流开源方案的全面调研，以及对2023–2026年最新研究进展的梳理，得出以下结论：'
    )
    bullets = [
        'ReFinED是目前综合性能最优的开源方案，在8个标准数据集上平均F1达89.4%，'
        '且推理效率高（比BLINK快6倍），适合作为基线参考',
        'LLM驱动的方法（LLMAEL、LELA）代表了最新趋势，在零样本设置下表现优异，'
        '但推理成本高，不适合大规模数据治理场景',
        'NIL检测仍是行业痛点，现有方案普遍表现不佳，是值得重点攻关的方向',
        '中文实体链接研究相对滞后，现有开源工具多以英文为主，中文支持有限',
        '全程可追溯的链接依据留痕是本智能体的独特价值，现有方案均未覆盖',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('6.2 技术路线建议', level=2)
    doc.add_paragraph('基于调研结果，建议本智能体的技术路线如下：')
    bullets = [
        '候选召回层：采用BM25全文检索（别名精确匹配）+ 向量语义检索（FAISS + m3e-base）的多路召回策略',
        '消歧排序层：基于实体类型过滤 + 名称最短编辑距离 + 上下文相似度的分层消歧',
        'NIL检测层：多信号融合（名称匹配度阈值 + 候选得分差距 + 类型一致性）',
        '追溯日志层：记录每一次链接的原值、新值、依据和置信度',
        '服务化层：FastAPI标准接口，支持单条/批量链接、追溯查询、知识库统计',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('6.3 差异化优势总结', level=2)
    doc.add_paragraph(
        '本智能体的核心研究价值不在于"复现已有开源能力"，而在于：'
    )
    bullets = [
        '将通用实体链接能力升级为面向治理、可控、可追溯、可集成的智能体',
        '在NIL检测、行业知识库适配、中文支持等"难点"上有所推进',
        '遵循"弱自主规划、强流程约束"的设计理念，契合数据治理对可审计的要求',
        '能用规则不用LLM，降低算力成本，提升可解释性和可复现性',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ========== 参考文献 ==========
    doc.add_heading('参考文献', level=1)

    refs = [
        '[1] Wu L, Petroni F, Josifoski M, et al. Scalable Zero-shot Entity Linking with Dense Entity Retrieval. EMNLP, 2020. (BLINK)',
        '[2] Ayoola T, Fisher S, Ponting J L, et al. An Efficient Zero-shot-capable Approach to End-to-End Entity Linking. NAACL Industry, 2022. (ReFinED)',
        '[3] De Cao N, Izacard G, Riedel S, et al. Autoregressive Entity Retrieval. ICLR, 2021. (GENRE)',
        '[4] van Hulst J, Hasibi F, Dercksen K, et al. REL: An Entity Linker Standing on the Shoulders of Giants. SIGIR, 2020.',
        '[5] Mendes P N, Jakob M, Garcia-Silva A, et al. DBpedia Spotlight: Shedding Light on the Web of Documents. I-Semantics, 2011.',
        '[6] Piccinno F, Ferragina P. From TagMe to WAT: a New Entity Annotator. ERD Workshop, 2014.',
        '[7] Usbeck R, Röder M, Ngonga Ngomo A C, et al. GERBIL – General Entity Annotator Benchmarking Framework. WWW, 2015.',
        '[8] Hofstätter S, Chen J, Raman K, et al. A Fair and In-Depth Evaluation of Existing End-to-End Entity Linking Systems. EMNLP, 2023.',
        '[9] Li Y, Galimov A, Ganapaneni M D, et al. ARTER: Adaptive Routing and Targeted Entity Reasoning. EMNLP, 2025. (Apple)',
        '[10] LLMAEL: Large Language Models are Good Context Augmenters for Entity Linking. CIKM, 2025. (清华KEG)',
        '[11] LELA: an LLM-based Entity Linking Approach with Zero-Shot Capability. arXiv:2601.05192, 2025.',
        '[12] FELA: Flexible Entity Linking Approach. CEUR-WS, 2025.',
        '[13] Xu Z F, Xin X. 基于大语言模型的中文实体链接实证研究. 自动化学报, 2025.',
        '[14] 何展鹏等. 基于知识库标记预训练孪生神经网络的中文实体链接. CSA, 2022.',
        '[15] 孙泽群. 大模型时代的知识融合. CCKS, 2024.',
        '[16] 郭浩等. 自适应特征融合的多模态实体对齐研究. 自动化学报, 2024.',
        '[17] Wang Z, Chen X. DERA: Dense Entity Retrieval for Entity Alignment in Knowledge Graphs. arXiv, 2024.',
        '[18] ELEVATE-ID: Efficient Uncertainty Estimation for LLM-based Entity Linking. DICE Research, 2025.',
        '[19] SynEL: A Synthetic Benchmark for Entity Linking. PMC, 2026.',
        '[20] Word Sense Disambiguation with Wikipedia Entities: A Survey of Entity Linking Approaches. PMC, 2025.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.runs[0].font.size = Pt(10)

    # ========== 保存 ==========
    output_path = 'F:/PyCharmProject/ELAGENT/data/同类开源方案调研报告.docx'
    doc.save(output_path)
    print(f'报告已保存至: {output_path}')
    return output_path


if __name__ == '__main__':
    create_report()
